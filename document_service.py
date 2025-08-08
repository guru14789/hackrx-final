import aiohttp
import asyncio
from typing import Dict, Any, List
import PyPDF2
import docx
import io
import re
import logging
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

class DocumentService:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
    
    async def process_document(self, document_url: str) -> Dict[str, Any]:
        """Process document from URL and return structured content"""
        try:
            logger.info(f"Processing document: {document_url}")
            
            # Download document
            document_data = await self._download_document(document_url)
            
            # Determine file type and extract content
            content_type = document_data["content_type"]
            file_content = document_data["content"]
            
            if "pdf" in content_type.lower():
                extracted_content = await self._extract_pdf_content(file_content)
            elif "word" in content_type.lower() or "docx" in content_type.lower():
                extracted_content = await self._extract_docx_content(file_content)
            else:
                extracted_content = await self._extract_text_content(file_content)
            
            # Create text chunks for embedding
            chunks = self._create_text_chunks(extracted_content["full_text"])
            
            result = {
                "content_type": content_type,
                "file_size": len(file_content),
                "page_count": extracted_content.get("page_count", 0),
                "full_text": extracted_content["full_text"],
                "chunks": chunks,
                "metadata": extracted_content.get("metadata", {}),
                "processing_info": {
                    "total_chunks": len(chunks),
                    "avg_chunk_size": sum(len(chunk) for chunk in chunks) // len(chunks) if chunks else 0
                }
            }
            
            logger.info(f"Document processed successfully: {len(chunks)} chunks created")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")
    
    async def _download_document(self, url: str) -> Dict[str, Any]:
        """Download document from URL"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=60)) as response:
                    if response.status != 200:
                        raise Exception(f"Failed to download document: HTTP {response.status}")
                    
                    content_type = response.headers.get('content-type', 'application/octet-stream')
                    content_length = int(response.headers.get('content-length', 0))
                    
                    if content_length > self.max_file_size:
                        raise Exception(f"File too large: {content_length} bytes (max: {self.max_file_size})")
                    
                    content = await response.read()
                    
                    return {
                        "content": content,
                        "content_type": content_type,
                        "size": len(content)
                    }
                    
        except Exception as e:
            logger.error(f"Error downloading document: {str(e)}")
            raise
    
    async def _extract_pdf_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            full_text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    cleaned_text = self._clean_text(page_text)
                    text_content.append({
                        'page': page_num + 1,
                        'content': cleaned_text
                    })
                    full_text_parts.append(cleaned_text)
            
            return {
                'text_content': text_content,
                'full_text': ' '.join(full_text_parts),
                'page_count': len(pdf_reader.pages),
                'metadata': {
                    'document_type': 'PDF',
                    'total_pages': len(pdf_reader.pages),
                    'pages_with_content': len(text_content)
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise Exception(f"Failed to extract PDF content: {str(e)}")
    
    async def _extract_docx_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from DOCX"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_content = []
            full_text_parts = []
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    cleaned_text = self._clean_text(paragraph.text)
                    text_content.append({
                        'paragraph': para_num + 1,
                        'content': cleaned_text
                    })
                    full_text_parts.append(cleaned_text)
            
            return {
                'text_content': text_content,
                'full_text': ' '.join(full_text_parts),
                'page_count': 1,  # DOCX doesn't have clear page breaks
                'metadata': {
                    'document_type': 'DOCX',
                    'total_paragraphs': len(doc.paragraphs),
                    'paragraphs_with_content': len(text_content)
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise Exception(f"Failed to extract DOCX content: {str(e)}")
    
    async def _extract_text_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from plain text"""
        try:
            text = content.decode('utf-8', errors='ignore')
            cleaned_text = self._clean_text(text)
            
            return {
                'text_content': [{'section': 1, 'content': cleaned_text}],
                'full_text': cleaned_text,
                'page_count': 1,
                'metadata': {'document_type': 'TEXT'}
            }
            
        except Exception as e:
            logger.error(f"Error extracting text content: {str(e)}")
            raise Exception(f"Failed to extract text content: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-$$$$\[\]\"\'\/\%\$]', '', text)
        
        # Fix common OCR errors
        text = text.replace('|', 'I')  # Common OCR mistake
        text = text.replace('0', 'O')  # In some contexts
        
        return text.strip()
    
    def _create_text_chunks(self, text: str) -> List[str]:
        """Create overlapping text chunks for better retrieval"""
        if not text:
            return []
        
        # Split by sentences first
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) + 1 <= self.chunk_size:
                current_chunk += sentence + ". "
            else:
                # Save current chunk if it has content
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                if len(chunks) > 0 and self.chunk_overlap > 0:
                    # Add some sentences from the end of previous chunk
                    prev_sentences = chunks[-1].split('. ')
                    overlap_sentences = prev_sentences[-2:] if len(prev_sentences) > 2 else prev_sentences
                    current_chunk = '. '.join(overlap_sentences) + ". " + sentence + ". "
                else:
                    current_chunk = sentence + ". "
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # Filter out very short chunks
        chunks = [chunk for chunk in chunks if len(chunk) > 50]
        
        logger.info(f"Created {len(chunks)} text chunks")
        return chunks
    
    def detect_document_domain(self, text: str) -> str:
        """Detect document domain based on content"""
        text_lower = text.lower()
        
        # Insurance keywords
        insurance_keywords = ['policy', 'premium', 'coverage', 'claim', 'deductible', 'beneficiary', 'insured']
        
        # Legal keywords
        legal_keywords = ['contract', 'agreement', 'clause', 'liability', 'jurisdiction', 'breach', 'damages']
        
        # HR keywords
        hr_keywords = ['employee', 'salary', 'benefits', 'leave', 'performance', 'termination', 'job description']
        
        # Compliance keywords
        compliance_keywords = ['regulation', 'compliance', 'audit', 'requirement', 'standard', 'procedure']
        
        # Count keyword occurrences
        scores = {
            'insurance': sum(1 for keyword in insurance_keywords if keyword in text_lower),
            'legal': sum(1 for keyword in legal_keywords if keyword in text_lower),
            'hr': sum(1 for keyword in hr_keywords if keyword in text_lower),
            'compliance': sum(1 for keyword in compliance_keywords if keyword in text_lower)
        }
        
        # Return domain with highest score
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        else:
            return 'general'
