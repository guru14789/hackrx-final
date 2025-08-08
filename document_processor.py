import requests
import PyPDF2
import docx
import io
import re
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document extraction and preprocessing"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def extract_content(self, document_url: str) -> Dict[str, Any]:
        """Extract content from various document formats"""
        try:
            # Download document
            response = requests.get(document_url, timeout=30)
            response.raise_for_status()
            
            # Determine file type
            content_type = response.headers.get('content-type', '').lower()
            
            if 'pdf' in content_type or document_url.lower().endswith('.pdf'):
                return self._extract_pdf_content(response.content)
            elif 'word' in content_type or document_url.lower().endswith(('.docx', '.doc')):
                return self._extract_docx_content(response.content)
            else:
                # Try to extract as text
                return self._extract_text_content(response.content)
                
        except Exception as e:
            logger.error(f"Error extracting content: {str(e)}")
            raise Exception(f"Failed to extract document content: {str(e)}")
    
    def _extract_pdf_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            metadata = {
                'total_pages': len(pdf_reader.pages),
                'document_type': 'PDF'
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append({
                        'page': page_num + 1,
                        'content': page_text.strip()
                    })
            
            return {
                'text_content': text_content,
                'metadata': metadata,
                'full_text': ' '.join([page['content'] for page in text_content])
            }
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise Exception(f"Failed to extract PDF content: {str(e)}")
    
    def _extract_docx_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from DOCX"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_content = []
            metadata = {
                'total_paragraphs': len(doc.paragraphs),
                'document_type': 'DOCX'
            }
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append({
                        'paragraph': para_num + 1,
                        'content': paragraph.text.strip()
                    })
            
            return {
                'text_content': text_content,
                'metadata': metadata,
                'full_text': ' '.join([para['content'] for para in text_content])
            }
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise Exception(f"Failed to extract DOCX content: {str(e)}")
    
    def _extract_text_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from plain text"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            return {
                'text_content': [{'section': 1, 'content': text}],
                'metadata': {'document_type': 'TEXT'},
                'full_text': text
            }
            
        except Exception as e:
            logger.error(f"Error extracting text content: {str(e)}")
            raise Exception(f"Failed to extract text content: {str(e)}")
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-$$$$]', '', text)
        
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks for better retrieval"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
            
            if i + chunk_size >= len(words):
                break
        
        return chunks
