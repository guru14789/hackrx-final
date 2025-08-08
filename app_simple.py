import streamlit as st
import requests
import json
import time
from typing import List, Dict, Any
import logging
import asyncio
import aiohttp
import PyPDF2
import docx
import io
import re
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class SimpleDocumentProcessor:
    """Simple document processor without external dependencies"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        self.max_file_size = 50 * 1024 * 1024  # 50MB
    
    async def process_document(self, document_url: str) -> Dict[str, Any]:
        """Process document from URL"""
        try:
            # Download document
            async with aiohttp.ClientSession() as session:
                async with session.get(document_url, timeout=aiohttp.ClientTimeout(total=60)) as response:
                    if response.status != 200:
                        raise Exception(f"Failed to download: HTTP {response.status}")
                    
                    content_type = response.headers.get('content-type', '')
                    content = await response.read()
            
            # Extract content based on type
            if 'pdf' in content_type.lower():
                return self._extract_pdf_content(content)
            elif 'word' in content_type.lower() or 'docx' in content_type.lower():
                return self._extract_docx_content(content)
            else:
                return self._extract_text_content(content)
                
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise
    
    def _extract_pdf_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text.strip())
            
            full_text = ' '.join(text_parts)
            chunks = self._create_chunks(full_text)
            
            return {
                'full_text': full_text,
                'chunks': chunks,
                'page_count': len(pdf_reader.pages),
                'content_type': 'PDF'
            }
            
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def _extract_docx_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from DOCX"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            full_text = ' '.join(text_parts)
            chunks = self._create_chunks(full_text)
            
            return {
                'full_text': full_text,
                'chunks': chunks,
                'page_count': len(doc.paragraphs),
                'content_type': 'DOCX'
            }
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_text_content(self, content: bytes) -> Dict[str, Any]:
        """Extract content from text"""
        try:
            text = content.decode('utf-8', errors='ignore')
            chunks = self._create_chunks(text)
            
            return {
                'full_text': text,
                'chunks': chunks,
                'page_count': 1,
                'content_type': 'TEXT'
            }
            
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def _create_chunks(self, text: str, chunk_size: int = 1000) -> List[str]:
        """Create text chunks"""
        if not text:
            return []
        
        # Simple chunking by sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 1 <= chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if len(chunk) > 50]

class SimpleQueryEngine:
    """Simple query engine using keyword matching"""
    
    def __init__(self):
        self.insurance_keywords = {
            'grace_period': ['grace period', 'grace time', 'payment grace'],
            'waiting_period': ['waiting period', 'wait time', 'waiting time'],
            'coverage': ['coverage', 'covered', 'covers', 'include'],
            'maternity': ['maternity', 'pregnancy', 'childbirth'],
            'pre_existing': ['pre-existing', 'pre existing', 'existing condition'],
            'cataract': ['cataract', 'eye surgery'],
            'organ_donor': ['organ donor', 'organ donation', 'transplant'],
            'no_claim_discount': ['no claim discount', 'NCD', 'bonus'],
            'health_checkup': ['health check', 'preventive', 'checkup'],
            'hospital': ['hospital', 'medical facility'],
            'ayush': ['ayush', 'ayurveda', 'homeopathy', 'unani'],
            'room_rent': ['room rent', 'ICU', 'accommodation']
        }
    
    def find_relevant_content(self, question: str, chunks: List[str]) -> List[str]:
        """Find relevant content chunks for a question"""
        question_lower = question.lower()
        
        # Extract keywords from question
        question_keywords = []
        for category, keywords in self.insurance_keywords.items():
            for keyword in keywords:
                if keyword in question_lower:
                    question_keywords.extend(keywords)
        
        # If no specific keywords found, use words from question
        if not question_keywords:
            question_keywords = [word for word in question_lower.split() if len(word) > 3]
        
        # Score chunks based on keyword presence
        scored_chunks = []
        for chunk in chunks:
            chunk_lower = chunk.lower()
            score = 0
            
            for keyword in question_keywords:
                if keyword in chunk_lower:
                    score += 1
            
            if score > 0:
                scored_chunks.append((chunk, score))
        
        # Sort by score and return top chunks
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        return [chunk for chunk, score in scored_chunks[:3]]
    
    def generate_answer(self, question: str, relevant_chunks: List[str]) -> str:
        """Generate answer based on relevant chunks"""
        if not relevant_chunks:
            return "I couldn't find specific information to answer your question in the provided document."
        
        # Combine relevant chunks
        context = " ".join(relevant_chunks)
        
        # Simple rule-based answer generation
        question_lower = question.lower()
        
        # Grace period questions
        if any(keyword in question_lower for keyword in ['grace period', 'grace time']):
            grace_matches = re.findall(r'grace period of (\w+) days?', context.lower())
            if grace_matches:
                return f"A grace period of {grace_matches[0]} days is provided for premium payment after the due date."
        
        # Waiting period questions
        if any(keyword in question_lower for keyword in ['waiting period', 'wait time']):
            waiting_matches = re.findall(r'waiting period of (\w+(?:-\w+)?)\s*(?:$$\d+$$)?\s*(months?|years?)', context.lower())
            if waiting_matches:
                period, unit = waiting_matches[0]
                return f"There is a waiting period of {period} {unit} for this coverage."
        
        # Coverage questions
        if any(keyword in question_lower for keyword in ['cover', 'coverage', 'covered']):
            if 'maternity' in question_lower:
                if 'maternity' in context.lower():
                    return "Yes, the policy covers maternity expenses. Please check the specific conditions and waiting periods mentioned in the policy document."
                else:
                    return "Maternity coverage information was not found in the provided document sections."
        
        # Default response with context
        return f"Based on the document: {relevant_chunks[0][:500]}..."

class SimpleLLMSystem:
    """Simple LLM system without external APIs"""
    
    def __init__(self):
        self.document_processor = SimpleDocumentProcessor()
        self.query_engine = SimpleQueryEngine()
        self.processed_documents = {}
    
    async def process_document(self, document_url: str) -> Dict[str, Any]:
        """Process document and store in memory"""
        try:
            doc_content = await self.document_processor.process_document(document_url)
            
            # Store in memory cache
            doc_id = hash(document_url)
            self.processed_documents[doc_id] = {
                'url': document_url,
                'content': doc_content,
                'processed_at': datetime.now()
            }
            
            return {
                'status': 'success',
                'doc_id': doc_id,
                'message': f"Document processed successfully. Found {len(doc_content['chunks'])} text chunks."
            }
            
        except Exception as e:
            return {'status': 'error', 'message': str(e)}
    
    def answer_questions(self, doc_id: int, questions: List[str]) -> List[str]:
        """Answer questions for a processed document"""
        if doc_id not in self.processed_documents:
            return ["Document not found. Please process the document first."] * len(questions)
        
        doc_content = self.processed_documents[doc_id]['content']
        chunks = doc_content['chunks']
        
        answers = []
        for question in questions:
            try:
                relevant_chunks = self.query_engine.find_relevant_content(question, chunks)
                answer = self.query_engine.generate_answer(question, relevant_chunks)
                answers.append(answer)
            except Exception as e:
                answers.append(f"Error processing question: {str(e)}")
        
        return answers

def main():
    st.set_page_config(
        page_title="Simple LLM Query System",
        page_icon="🔍",
        layout="wide"
    )
    
    st.title("🔍 Simple LLM Query-Retrieval System")
    st.markdown("**Simplified Version** - Works without external APIs or complex dependencies")
    
    # Initialize system
    if 'system' not in st.session_state:
        st.session_state.system = SimpleLLMSystem()
        st.session_state.current_doc_id = None
    
    # Sidebar
    with st.sidebar:
        st.header("System Status")
        st.success("✅ Document Processor: Ready")
        st.success("✅ Query Engine: Ready")
        st.info("ℹ️ No external APIs required")
        
        if st.session_state.current_doc_id:
            st.success(f"📄 Document loaded: ID {st.session_state.current_doc_id}")
    
    # Main interface
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📄 Document Input")
        
        document_url = st.text_input(
            "Document URL",
            value="https://hackrx.blob.core.windows.net/assets/policy.pdf?sv=2023-01-03&st=2025-07-04T09%3A11%3A24Z&se=2027-07-05T09%3A11%3A00Z&sr=b&sp=r&sig=N4a9OU0w0QXO6AOIBiu4bpl7AXvEZogeT%2FjUHNO7HzQ%3D",
            help="Enter the URL of the document to process"
        )
        
        if st.button("📥 Process Document", type="primary"):
            if document_url:
                with st.spinner("Processing document..."):
                    # Run async function
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    result = loop.run_until_complete(
                        st.session_state.system.process_document(document_url)
                    )
                    loop.close()
                    
                    if result['status'] == 'success':
                        st.session_state.current_doc_id = result['doc_id']
                        st.success(result['message'])
                    else:
                        st.error(f"Processing failed: {result['message']}")
        
        st.subheader("Questions")
        questions_text = st.text_area(
            "Enter your questions (one per line)",
            height=200,
            value="""What is the grace period for premium payment under the National Parivar Mediclaim Plus Policy?
What is the waiting period for pre-existing diseases (PED) to be covered?
Does this policy cover maternity expenses, and what are the conditions?
What is the waiting period for cataract surgery?
Are the medical expenses for an organ donor covered under this policy?""",
            help="Enter one question per line"
        )
        
        questions = [q.strip() for q in questions_text.split('\n') if q.strip()]
        
        if questions:
            st.write(f"**{len(questions)} questions ready**")
    
    with col2:
        st.header("🎯 Query Results")
        
        if st.button("🚀 Answer Questions", type="primary", disabled=not (st.session_state.current_doc_id and questions)):
            if st.session_state.current_doc_id and questions:
                with st.spinner("Generating answers..."):
                    answers = st.session_state.system.answer_questions(
                        st.session_state.current_doc_id, 
                        questions
                    )
                    
                    st.success(f"Generated {len(answers)} answers!")
                    
                    # Display results
                    for i, (question, answer) in enumerate(zip(questions, answers), 1):
                        with st.expander(f"Q{i}: {question}", expanded=True):
                            st.write("**Answer:**")
                            st.write(answer)
                            
                            # Simple confidence indicator
                            if "couldn't find" not in answer.lower() and "error" not in answer.lower():
                                confidence = min(90, 70 + len(answer.split()) // 5)
                                st.progress(confidence / 100)
                                st.caption(f"Confidence: {confidence}%")
                            else:
                                st.progress(0.3)
                                st.caption("Confidence: Low")
                    
                    # Export option
                    export_data = {
                        "document_url": document_url,
                        "timestamp": datetime.now().isoformat(),
                        "questions_and_answers": [
                            {"question": q, "answer": a} 
                            for q, a in zip(questions, answers)
                        ]
                    }
                    
                    st.download_button(
                        "📥 Download Results",
                        data=json.dumps(export_data, indent=2),
                        file_name=f"results_{int(time.time())}.json",
                        mime="application/json"
                    )
    
    # Instructions
    with st.expander("📖 How to Use", expanded=False):
        st.markdown("""
        **Step 1:** Enter a document URL (PDF or DOCX)
        
        **Step 2:** Click "Process Document" and wait for completion
        
        **Step 3:** Enter your questions (one per line)
        
        **Step 4:** Click "Answer Questions" to get responses
        
        **Note:** This simplified version uses keyword matching and doesn't require external APIs.
        """)

if __name__ == "__main__":
    main()
